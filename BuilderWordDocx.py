import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import datetime

from numpy import datetime_as_string

EXT_FILE = ".docx"

class BuilderWordDocx():

    def __init__(self) -> None:
        self.doc = None
        self.paragraph_text = None
        
    def builder_name_file(self,nameFile="default")->str:
        if nameFile == None or nameFile == "":
            nameFile = "default"
        else:
            nameFile = nameFile.replace(" ","_")

        date_time = datetime.datetime.now()
        date_time_format = date_time.strftime("%Y-%m-%d-%X")  
        return nameFile +'-'+date_time_format.replace(':','-')

    def create_docx(self,nameFile,text)-> None:
        self.nameFile = self.builder_name_file(nameFile)
        self.doc = docx.Document()
        self.doc.add_heading(self.nameFile,1)
        self.paragraph_text = self.doc.add_paragraph(text)
        self.doc.save(self.return_file_name(self.path_save_document)) 
    
    def configure_path_save_document(self,path):
        self.path_save_document = path

    def return_file_name(self,path_save_document):
        return os.path.normpath(os.path.join(path_save_document,self.nameFile+EXT_FILE)) 
